from pathlib import Path
from tkinter import filedialog
from openpyxl import load_workbook
from unicodedata import normalize
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from shutil import copy2

import tkinter as tk

# Bloque de codigo para seleccionar donde se crear el Word
def seleccionar_carpeta_destino():
    ventana = tk.Tk()
    ventana.withdraw()
    ventana.attributes("-topmost", True)

    ruta = filedialog.askdirectory(
        title="Selecciona donde crear el Documento"
    )

    ventana.destroy()

    if not ruta:
        print("No se selecciono una carpeta")
        return None

    return Path(ruta)

carpeta_destino = seleccionar_carpeta_destino()

# Bloque de codigo para seleccionar el archivo a leer
def seleccionar_xslx():
    ventana = tk.Tk()
    ventana.withdraw()
    ventana.attributes("-topmost", True)

    archivo = filedialog.askopenfilename(
        title="Selecciona el archivo Excel a leer",
        filetypes=[
            ("Archivos Excel", "*.xlsx *.xls"),
            ("Todos los archivos", "*.*")
        ]
    )

    ventana.destroy()
    return Path(archivo) if archivo else None

# Modificacion del texto para evitar errores
def normalizar_texto(valor):
    # Converitr a mayusculas sin tildes
    texto = str(valor or "").strip().upper()
    return "".join(
        caracter
        for caracter in normalize("NFD", texto)
        if not ("\u0300" <= caracter <= "\u036F")
    )

# Obtener datos de los alumnos
FILA_INICIAL = 3
FILA_FINAL = 22
def leer_alumnos(archivo_xls):
    libro = load_workbook(archivo_xls, data_only=True)
    hoja = libro.worksheets[0]

    print(f"Hoja que se leerá: {hoja.title}")

    encabezados = {
        normalizar_texto(celda.value): celda.column
        for celda in hoja[1]
        if celda.value
    }

    columnas__necesarias = {
        "1 COGNOM": "apellido_1",
        "2 COGNOM": "apellido_2",
        "NOM": "nombre",
        "DNI/NIE": "nif",
    }

    for encabezado in columnas__necesarias:
        if encabezado not in encabezados:
            raise ValueError(
                f"No se encontró la columna «{encabezado}» en la fila 1."
            )

    alumnos = []

    for fila in hoja.iter_rows(min_row=FILA_INICIAL, max_row=FILA_FINAL, values_only=True):
        apellido_1 = fila[encabezados["1 COGNOM"] - 1]
        apellido_2 = fila[encabezados["2 COGNOM"] - 1]
        nombre = fila[encabezados["NOM"] - 1]
        nif = fila[encabezados["DNI/NIE"] - 1]

        # ignorar filas sin alumnos
        if not any([apellido_1,apellido_2,nombre,nif]):
            continue

        alumnos.append({
            "apellido_1": str(apellido_1 or "").strip(),
            "apellido_2": str(apellido_2 or "").strip(),
            "nombre": str(nombre or "").strip(),
            "nif": str(nif or "").strip(),
        })

    libro.close()
    return alumnos

# Ordenar por apellidos a los alumnos
def ordenar_alumnos(alumnos):
    return sorted(
        alumnos,
        key=lambda alumno:(
            normalizar_texto(alumno["apellido_1"]),
            normalizar_texto(alumno["apellido_2"]),
            normalizar_texto(alumno["nombre"]),
        )
    )

# Leer hoja 2 del excel - datos de control
def leer_datos_control(archivo_xls):
    libro = load_workbook(archivo_xls, data_only=True)

    # Segunda hoja
    hoja = libro.worksheets[1]

    print(f"Hoja de datos de control: {hoja.title}")

    datos_control = []

    for numero_fila, fila in enumerate(hoja.iter_rows(), start=1):
        # Leer celdas con contenido
        celdas_con_valor = [
            celda
            for celda in fila
            if celda.value is not None and str(celda.value).strip()
        ]

        # Ignora filas sin datos
        if len(celdas_con_valor) < 2:
            continue

        etiqueta = str(celdas_con_valor[0].value).strip()
        valor = str(celdas_con_valor[-1].value).strip()

        datos_control.append({
            "fila_excel": numero_fila,
            "etiqueta": etiqueta,
            "valor": valor,
        })

    libro.close()
    return datos_control

# Creacion del documento 
PLANTILLA = Path(__file__).with_name("plantilla_control.docx")
def obtener_dato_por_fila(datos_control, fila_excel):
    for dato in datos_control:
        if dato["fila_excel"] == fila_excel:
            return dato["valor"]

    return ""

def crear_documento_desde_plantilla(carpeta_destino, datos_control, alumnos):
    if not PLANTILLA.exists():
        raise FileNotFoundError(
            "No se encontró «plantilla_control.docx» junto a main.py."
        )

    numero_homologacion = obtener_dato_por_fila(datos_control, 2)
    fecha_inicio = obtener_dato_por_fila(datos_control, 10)

    fecha_archivo = fecha_inicio[:10].replace("-", "")
    nombre_archivo = f"{numero_homologacion}_{fecha_archivo}.docx"
    ruta_word = carpeta_destino / nombre_archivo

    # 1. Copiamos el diseño original.
    copy2(PLANTILLA, ruta_word)

    # 2. Abrimos la copia, nunca el archivo de plantilla.
    documento = Document(ruta_word)

    # La primera tabla contiene cabecera, curso y asistentes.
    tabla = documento.tables[0]
    filas = tabla.rows

    # --- Datos del centro ---
    escribir_etiqueta_valor(
        primer_parrafo_util(filas[1].cells[0]),
        "Número d’homologació del curs",
        numero_homologacion
    )

    escribir_etiqueta_valor(
        primer_parrafo_util(filas[2].cells[0]),
        "Número identificació del centre",
        obtener_dato_por_fila(datos_control, 3)
    )

    escribir_texto(
        primer_parrafo_util(filas[3].cells[0]),
        "Nom del centre",
        tamaño=8
    )
    escribir_texto(
        filas[3].cells[0].paragraphs[-1],
        obtener_dato_por_fila(datos_control, 4),
        tamaño=10
    )

    escribir_texto(
        primer_parrafo_util(filas[4].cells[0]),
        "Adreça postal",
        tamaño=8
    )
    escribir_texto(
        filas[4].cells[0].paragraphs[-1],
        obtener_dato_por_fila(datos_control, 5),
        tamaño=10
    )

    escribir_texto(
        primer_parrafo_util(filas[5].cells[0]),
        "Codi postal",
        tamaño=8
    )
    escribir_texto(
        filas[5].cells[0].paragraphs[-1],
        obtener_dato_por_fila(datos_control, 6),
        tamaño=10
    )

    escribir_texto(
        primer_parrafo_util(filas[5].cells[5]),
        "Localitat",
        tamaño=8
    )
    escribir_texto(
        filas[5].cells[5].paragraphs[-1],
        obtener_dato_por_fila(datos_control, 7),
        tamaño=10
    )

    escribir_texto(
        primer_parrafo_util(filas[5].cells[8]),
        "Telèfon",
        tamaño=8
    )
    escribir_texto(
        filas[5].cells[8].paragraphs[-1],
        obtener_dato_por_fila(datos_control, 8),
        tamaño=10
    )

    # --- Fechas y horario ---
    fecha_final = obtener_dato_por_fila(datos_control, 11)
    horario = obtener_dato_por_fila(datos_control, 12)

    escribir_texto(
        primer_parrafo_util(filas[10].cells[0]),
        f"     {fecha_inicio[:10].replace("-", "/")}        "
        f"Data de finalització   {fecha_final[:10].replace("-", "/")}",
        tamaño=8
    )

    escribir_texto(
        primer_parrafo_util(filas[11].cells[0]),
        f"Sessió dia     {fecha_inicio[:10].replace("-", "/")}",
        tamaño=8
    )

    escribir_texto(
        primer_parrafo_util(filas[11].cells[3]),
        f"Matí  de    {horario}",
        tamaño=8
    )

    # --- Alumnos: 10 en la primera página y 10 en la segunda ---
    filas_alumnos = list(range(14, 24)) + list(range(35, 45))

    for numero, fila_excel in enumerate(filas_alumnos, start=1):
        fila_word = filas[fila_excel]

        if numero <= len(alumnos):
            alumno = alumnos[numero - 1]

            apellidos = (
                f'{alumno["apellido_1"]} {alumno["apellido_2"]}'
            ).strip()

            nombre = alumno["nombre"]
            nif = alumno["nif"]
        else:
            apellidos = ""
            nombre = ""
            nif = ""

        escribir_texto(
            primer_parrafo_util(fila_word.cells[0]),
            f"{numero}.",
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[2]),
            apellidos,
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[6]),
            nombre,
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[8]),
            nif,
            negrita=True
        )

    # --- Docentes ---
    docentes = [
        {
            "nombre": obtener_dato_por_fila(datos_control, 13),
            "apellidos": obtener_dato_por_fila(datos_control, 14),
            "nif": obtener_dato_por_fila(datos_control, 15),
        },
        {
            "nombre": obtener_dato_por_fila(datos_control, 16),
            "apellidos": obtener_dato_por_fila(datos_control, 17),
            "nif": obtener_dato_por_fila(datos_control, 18),
        },
    ]

    # Docentes mostrados al final de la primera página.
    for indice, docente in enumerate(docentes):
        fila_word = filas[26 + indice]

        escribir_texto(
            primer_parrafo_util(fila_word.cells[0]),
            docente["apellidos"],
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[4]),
            docente["nombre"],
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[7]),
            docente["nif"],
            negrita=True
        )

    # Docentes repetidos en la segunda página.
    tabla_docentes = documento.tables[1]

    for indice, docente in enumerate(docentes):
        fila_word = tabla_docentes.rows[2 + indice]

        escribir_texto(
            primer_parrafo_util(fila_word.cells[0]),
            docente["apellidos"],
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[1]),
            docente["nombre"],
            negrita=True
        )
        escribir_texto(
            primer_parrafo_util(fila_word.cells[2]),
            docente["nif"],
            negrita=True
        )

    documento.save(ruta_word)

    print(f"\nDocumento creado correctamente:\n{ruta_word}")

def configurar_run(run, negrita=False, tamaño=10):
    run.font.name="Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(tamaño)
    run.bold = negrita

def primer_parrafo_util(celda):
    for parrafo in celda.paragraphs:
        if parrafo.text.strip():
            return parrafo

    return celda.paragraphs[-1]

def escribir_texto(parrafo, texto, negrita=False, tamaño=10):
    parrafo.clear()
    #Clear() conserva los detalles del parrafo

    run = parrafo.add_run(str(texto or ""))
    configurar_run(run, negrita, tamaño)

def escribir_etiqueta_valor(parrafo, etiqueta, valor):
    parrafo.clear()

    run_etiqueta = parrafo.add_run(f"{etiqueta}: ")
    configurar_run(run_etiqueta, negrita=True, tamaño=10)

    run_valor = parrafo.add_run(str(valor or ""))
    configurar_run(run_valor, negrita=False, tamaño=10)

def obtener_dato_por_fila(datos_control, fila_excel):
    for dato in datos_control:
        if dato["fila_excel"] == fila_excel:
            return dato["valor"]

    return ""

# Validar si las selecciones fueron correctas
if carpeta_destino is None:
    print("Operacion cancelada")
else:
    archivo_xls = seleccionar_xslx()

    if archivo_xls is None:
        print ("Operacion cancelada")

    else:
        print(f"El Word se guardara en: {carpeta_destino}")
        print(f"Archivo seleccionado: {archivo_xls.name}")

alumnos = leer_alumnos(archivo_xls)
alumnos = ordenar_alumnos(alumnos)
print(f"\nSe han leído {len(alumnos)} alumnos:\n")

datos_control = leer_datos_control(archivo_xls)
print("\nDATOS DEL CONTROL:\n")

crear_documento_desde_plantilla(
    carpeta_destino,
    datos_control,
    alumnos
)

for alumno in alumnos:
    print(alumno)

for dato in datos_control:
    print(f'{dato["etiqueta"]}: {dato["valor"]}')